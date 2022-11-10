import random
from docx import Document
from docx.shared import Pt
from main import first_level, second_level, third_level

document = Document()
equations = Document()
apple = b'\xf0\x9f\x8d\x8e'.decode("utf-8")
pear = b'\xf0\x9f\x8d\x90'.decode("utf-8")
lemon = b'\xf0\x9f\x8d\x8b'.decode("utf-8")
plus = b'\xE2\x9E\x95'.decode("utf-8")
minus = b'\xE2\x9E\x96'.decode("utf-8")
multiply = b'\xE2\x9C\x96'.decode("utf-8")
equals = b'\xEF\xBC\x9D'.decode("utf-8")


if first_level != 0:
    document.add_heading("Answers 1 level")
    p = document.add_paragraph()

for i in range(first_level):
    equation1 = equations.add_paragraph()
    equation1.style.font.size = Pt(20)
    equation1.paragraph_format.keep_together = True

    s = random.choice(list(open('first-pr.txt')))
    x = random.randint(5, 100)
    a = random.randint(6, 100)
    while a == x:
        a = random.randint(10, 100)

    if s[2] == "+":
        s = str(s).replace('a', str(a))
        s = str(s).replace('z', str(a + x))
    elif s[2] == "*":
        s = str(s).replace('a', str(a))
        s = str(s).replace('z', str(a * x))
    else:
        if s[0] == "a":
            s = str(s).replace('a', str(a))
            s = str(s).replace('z', str(a - x))
        else:
            s = str(s).replace('a', str(a))
            s = str(s).replace('z', str(x - a))

    s = s.replace("x", str(apple))
    s = s.replace("+", plus)
    s = s.replace("-", minus)
    s = s.replace(" ", "   ")
    s = s.replace("*", multiply)
    s = s.replace("=", equals)
    equation1.add_run(s)
    equation1.add_run("-----------------------------------------------------------------")

    p.add_run(str(x))
    if i != first_level-1:
        p.add_run("\n")


if second_level != 0:
    document.add_heading("Answers 2 level")
    p2 = document.add_paragraph()

for i in range(second_level):
    equation1 = equations.add_paragraph()
    equation1.style.font.size = Pt(20)
    equation1.paragraph_format.keep_together = True

    s1 = random.choice(list(open('second_pr.txt')))
    x = random.randint(5, 100)
    y = random.randint(10, 100)

    while y == x:
        y = random.randint(10, 100)

    if s1[2] == '+':
        if s1[11] == "y" or s1[15] == "y":
            if x >= y:
                s1 = str(s1).replace('z', str(x - y))
            else:
                s1 = str(s1).replace('z', "(" + str(x - y) + ")")
        else:
            if x < y:
                s1 = str(s1).replace('z', str(y - x))
            else:
                s1 = str(s1).replace('z', "(" + str(y - x) + ")")

        s1 = str(s1).replace('a', str(x + y))

    else:
        if s1[11] == "x" and s1[13] == "-":
            if x >= y:
                s1 = str(s1).replace('z', str(x - y))
            else:
                s1 = str(s1).replace('z', "(" + str(x - y) + ")")
        else:
            s1 = str(s1).replace('z', str(x + y))

        if x >= y:
            s1 = str(s1).replace('a', str(x - y))
        else:
            s1 = str(s1).replace('a', "(" + str(x - y) + ")")

    s1 = s1.replace("x", apple)
    s1 = s1.replace("y", lemon)
    s1 = s1.replace("+", plus)
    s1 = s1.replace("-", minus)
    s1 = s1.replace(" ", "   ")
    s1 = s1.replace("*", multiply)
    s1 = s1.replace("=", equals)

    equation1.add_run(s1[:s1.find(";")] + "\n")
    equation1.add_run(s1[s1.find(";")+4:])
    equation1.add_run("-----------------------------------------------------------------")

    p2.add_run(str(x) + "   " + str(y))
    if i != second_level-1:
        p2.add_run("\n")


if third_level != 0:
    document.add_heading("Answers 3 level")
    p3 = document.add_paragraph()

for i in range(third_level):
    equation1 = equations.add_paragraph()
    equation1.style.font.size = Pt(20)
    equation1.paragraph_format.keep_together = True

    s2 = random.choice(list(open('third_pr.txt')))
    x = random.randint(5, 100)
    y = random.randint(10, 100)
    z = random.randint(15, 100)
    if s2[2] == "+" and s2[6] == "+":
        if s2[17] == "+" or s2[17] == "=":
            s2 = str(s2).replace('a', str(x + y + z))
            if z >= y:
                s2 = str(s2).replace('b', str(z - y))
            else:
                s2 = str(s2).replace('b', "(" + str(z - y) + ")")
            s2 = str(s2).replace('c', str(z + x))
        else:
            s2 = str(s2).replace('a', str(x + y + z))
            if y >= z:
                s2 = str(s2).replace('b', str(y - z))
            else:
                s2 = str(s2).replace('b', "(" + str(y - z) + ")")
            s2 = str(s2).replace('c', str(z + x))

    elif s2[2] == "+" and s2[6] == "-":
        if s2[17] == "+" or s2[17] == "=":
            if x + y - z >= 0:
                s2 = str(s2).replace('a', str(x + y - z))
            else:
                s2 = str(s2).replace('a', "(" + str(x + y - z) + ")")

            if z >= y:
                s2 = str(s2).replace('b', str(z - y))
            else:
                s2 = str(s2).replace('b', "(" + str(z - y) + ")")
            s2 = str(s2).replace('c', str(z + x))

        else:
            if x + y - z >= 0:
                s2 = str(s2).replace('a', str(x + y - z))
            else:
                s2 = str(s2).replace('a', "(" + str(x + y - z) + ")")

            if y >= z:
                s2 = str(s2).replace('b', str(y - z))
            else:
                s2 = str(s2).replace('b', "(" + str(y - z) + ")")

            s2 = str(s2).replace('c', str(z + x))

    elif s2[2] == "-" and s2[6] == "-":
        if s2[17] == "+" or s2[17] == "=":
            if x - y - z >= 0:
                s2 = str(s2).replace('a', str(x - y - z))
            else:
                s2 = str(s2).replace('a', "(" + str(x - y - z) + ")")

            if z >= y:
                s2 = str(s2).replace('b', str(z - y))
            else:
                s2 = str(s2).replace('b', "(" + str(z - y) + ")")

            s2 = str(s2).replace('c', str(z + x))

        else:
            if x - y - z >= 0:
                s2 = str(s2).replace('a', str(x - y - z))
            else:
                s2 = str(s2).replace('a', "(" + str(x - y - z) + ")")

            if y >= z:
                s2 = str(s2).replace('b', str(y - z))
            else:
                s2 = str(s2).replace('b', "(" + str(y - z) + ")")

            s2 = str(s2).replace('c', str(z + x))

    s2 = s2.replace("x", apple)
    s2 = s2.replace("y", lemon)
    s2 = s2.replace("z", pear)
    s2 = s2.replace("+", plus)
    s2 = s2.replace("-", minus)
    s2 = s2.replace(" ", "   ")
    s2 = s2.replace("*", multiply)
    s2 = s2.replace("=", equals)

    equation1.add_run(s2[:s2.find(";")] + "\n")
    equation1.add_run(s2[s2.find(";")+4:s2.find(";", s2.find(";") + 1)] + "\n")
    equation1.add_run(s2[s2.find(";", s2.find(";") + 1)+4:])
    equation1.add_run("-----------------------------------------------------------------")

    p3.add_run(str(x) + "   " + str(y) + "   " + str(z))
    if i != third_level-1:
        p3.add_run("\n")


document.save("Answers.docx")
equations.save("Equations.docx")
